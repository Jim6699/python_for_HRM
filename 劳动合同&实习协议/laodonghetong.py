from docx import Document
from openpyxl import load_workbook
import os
import datetime
import time
# 结合路径判断生成文件夹，规避程序报错而终止的风险
if not os.path.exists('./' + '全部合同'):
    os.mkdir('./' + '全部合同')
    print("创建目录成功")
 
workbook = load_workbook( './' + '批量合同.xlsx')
sheet = workbook.active
#人数
numbers = int(sheet.max_row)-1
print("人数共计：{}人".format(str(numbers)))
 
for table_row in range(2, sheet.max_row + 1):
    wordfile = Document('./' + '劳动合同【模板勿动】.docx')
    print("=====合同姓名类型=====")
    name = str(sheet.cell(row=table_row, column=1).value)
    print(name)
    series = str(sheet.cell(row=table_row, column=6).value)
    print(series)
    try:
        #获取合同开始日期
        print("=====合同开始日期=====")
        date_kaishi =  sheet['B{}'.format(table_row)].value
        try:
            date_kaishi_year = "20"+str(date_kaishi.strftime('%y')) 
            date_kaishi_month = str(date_kaishi.strftime('%m'))
            date_kaishi_day = str(date_kaishi.strftime('%d'))
        except:
            date_kaishi_year = datetime.datetime.strptime(date_kaishi,'%Y-%m-%d').year
            date_kaishi_month = datetime.datetime.strptime(date_kaishi,'%Y-%m-%d').month
            date_kaishi_day = datetime.datetime.strptime(date_kaishi,'%Y-%m-%d').day
        print(date_kaishi_year)
        print(date_kaishi_month)
        print(date_kaishi_day)
        #获取合同终止日期
        print("=====合同终止日期=====")
        date_zhongzhi = sheet['C{}'.format(table_row)].value
        try:
            date_zhongzhi_year = "20"+str(date_zhongzhi.strftime('%y')) 
            date_zhongzhi_month = str(date_zhongzhi.strftime('%m'))
            date_zhongzhi_day = str(date_zhongzhi.strftime('%d'))
        except:
            date_zhongzhi_year = datetime.datetime.strptime(date_zhongzhi,'%Y-%m-%d').year
            date_zhongzhi_month = datetime.datetime.strptime(date_zhongzhi,'%Y-%m-%d').month
            date_zhongzhi_day = datetime.datetime.strptime(date_zhongzhi,'%Y-%m-%d').day
        print(date_zhongzhi_year)
        print(date_zhongzhi_month)
        print(date_zhongzhi_day)
        #获取合同试用日期
        print("=====试用截止日期=====")
        date_shiyong = sheet['D{}'.format(table_row)].value
        print(date_shiyong)
        if date_shiyong ==  None:
            print("无试用期")
            date_shiyong_if = "/"
            date_shiyong_year = "/"
            date_shiyong_month = "/"
            date_shiyong_day = "/"
            date_shiyong_year_1 = "/"
            date_shiyong_month_1 = "/"
            date_shiyong_day_1 = "/"
        else:
            date_shiyong_year = date_kaishi_year
            date_shiyong_month = date_kaishi_month
            date_shiyong_day = date_kaishi_day
            try:
                date_shiyong_if = "陆"
                date_shiyong_year_1 = "20"+str(date_shiyong.strftime('%y')) 
                date_shiyong_month_1 = str(date_shiyong.strftime('%m'))
                date_shiyong_day_1 = str(date_shiyong.strftime('%d'))
            except:
                date_shiyong_if = "陆"
                date_shiyong_year_1 = datetime.datetime.strptime(date_shiyong,'%Y-%m-%d').year
                date_shiyong_month_1 = datetime.datetime.strptime(date_shiyong,'%Y-%m-%d').month
                date_shiyong_day_1 = datetime.datetime.strptime(date_shiyong,'%Y-%m-%d').day
        print(date_shiyong_if)
        print(date_shiyong_year)
        print(date_shiyong_month)
        print(date_shiyong_day)        
        print(date_shiyong_year_1)
        print(date_shiyong_month_1)
        print(date_shiyong_day_1)
        #获取合同金额
        print("=====试用合同金额=====")
        money = sheet['E{}'.format(table_row)].value
        print(money)
 
        if series == "劳动合同":
            wordfile = Document('./' + '劳动合同【模板勿动】.docx')
        elif series == "实习协议":
            wordfile = Document('./' + '实习协议【模板勿动】.docx')
        else:
            print("暂时不支持本合同类型，跳过！")
            continue
        all_paragraphs = wordfile.paragraphs
        for paragraph in all_paragraphs:
            for run in paragraph.runs:
                try:
                    if "奰" in run.text:
                        run.text = run.text.replace("奰", date_kaishi_year)
                    elif "躄" in run.text:  
                        run.text = run.text.replace("躄", date_kaishi_month)
                    elif "罍" in run.text:  
                        run.text = run.text.replace("罍", date_kaishi_day)
                    elif "颣" in run.text:  
                        run.text = run.text.replace("颣", date_zhongzhi_year)
                    elif "薐" in run.text:   
                        run.text = run.text.replace("薐", date_zhongzhi_month)
                    elif "豳" in run.text:
                        run.text = run.text.replace("豳", date_zhongzhi_day)
                    elif "懿" in run.text:
                        run.text = run.text.replace("懿", str(money))
                    elif "鰘" in run.text:
                        run.text = run.text.replace("鰘", date_shiyong_if)
                    elif "翳" in run.text:
                        run.text = run.text.replace("翳", date_shiyong_year)
                    elif "薹" in run.text: 
                        run.text = run.text.replace("薹", date_shiyong_month)
                    elif "虩" in run.text: 
                        run.text = run.text.replace("虩", date_shiyong_day)
                    elif "舄" in run.text:  
                        run.text = run.text.replace("舄", date_shiyong_year_1)
                    elif "衚" in run.text:
                        run.text = run.text.replace("衚", date_shiyong_month_1)
                    elif "衕" in run.text:
                        run.text = run.text.replace("衕", date_shiyong_day_1)                    
                    
                except Exception as e:
                    print("替换文本出错："+str(e)) 
            
    except Exception as e:
        print("出错："+str(e)) 
    wordfile.save('./' + f'全部合同/{name}_{series}.docx')
    print(f"{name}_{series}.docx | 另存成功")
input ("Please Enter to close this exe:")
